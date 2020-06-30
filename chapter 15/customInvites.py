import docx,os

os.chdir(r'C:\Users\YourName\AppData\Local\Programs\Python\Python38-32\atbswp\practiceProg')
myfile=open('guests.txt','r')
names=myfile.readlines()

document=docx.Document('invite.docx')
for name in names: # same repeates for each name on same doc but seperate pages
    name=name.strip() # to make the invite tidy incase our guests list has unnecessary spaces
    # all these styles have been created on the invite.docx and saved blank

    document.add_paragraph('It would be a pleasure to have the company of ',style='Style1')
    document.add_paragraph(name,style='Style2')
    document.add_paragraph('at 11010 Memory Lane on the Evening of ',style='Style1')
    document.add_paragraph('April 1st',style='Style3')
    document.add_paragraph("at 7'o clock",style='Style1')
    document.add_page_break() # each new invite on new page
    
document.save('invites.docx') # saved it in another document

print('Customized Invites have been saved in invites.docx file in same folder')

